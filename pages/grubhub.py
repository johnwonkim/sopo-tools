from docx import Document
from docx.table import Table
from docx.oxml.table import CT_Tbl
import copy
import re
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING
import email
from email import policy
from bs4 import BeautifulSoup
import streamlit as st
import io

TEMPLATE_PATH = "pages/templates/labels_template.docx"
LABELS_PER_PAGE = 10

def convert_eml_to_html(file_contents):
    # Read raw email
    msg = email.message_from_bytes(file_contents, policy=policy.default)
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == "text/html":
                decoded_body = part.get_content()
            elif content_type == "text/plain":
                decoded_body = part.get_content()
    else:
        decoded_body = msg.get_content()

    return BeautifulSoup(decoded_body, "html.parser")

# Extract order date and delivery name
def extract_order_info(soup):
    """
    Extracts the order date and delivery name from a BeautifulSoup object.

    Args:
        soup (BeautifulSoup): Parsed HTML content.

    Returns:
        (order_time, delivery_name): Tuple of strings or None if not found.
    """
    order_time = None
    delivery_name = None

    # ---- Extract Order Date ----
    date_label = soup.find(string=lambda text: text and "Order placed on:" in text)
    if date_label:
        try:
            date_span = date_label.find_parent().find_next("span")
            order_time = date_span.get_text(strip=True)
        except Exception:
            pass

    # Fallback: regex
    if not order_time:
        text = soup.get_text()
        match = re.search(
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)'
            r'\s+\d{1,2},\s+\d{4},\s+\d{1,2}:\d{2}:\d{2}\s+(?:AM|PM)',
            text
        )
        if match:
            order_time = match.group()

    # ---- Extract Deliver To ----
    deliver_label = soup.find(string=lambda text: text and "Deliver to:" in text)
    if deliver_label:
        try:
            name_div = deliver_label.find_parent().find_next("div")
            delivery_name = name_div.get_text(strip=True)
        except Exception:
            pass

    return order_time, delivery_name

def extract_individual_labels_from_html(soup):
    labels = []

    order_date, delivery_name = extract_order_info(soup)

    customer_blocks = soup.find_all("table", style=lambda s: s and "border-bottom:dashed thin black" in s)

    for block in customer_blocks:
        try:
            name_block = block.find("span", string=re.compile(r"\d+\s*/\s*\d+"))
            if not name_block:
                continue

            order_id = name_block.get_text(strip=True)
            name_span = name_block.find_previous("span")
            name = name_span.get_text(strip=True) if name_span else "Unknown"
            name_line = f"{name} {order_id}"

            items_table = block.find("table", style=lambda s: s and "width:390px" in s)
            if not items_table:
                continue

            tbody = items_table.find("tbody")
            if not tbody:
                continue

            rows = tbody.find_all("tr")
            items = []

            for row in rows:
                cols = row.find_all("td")
                if len(cols) == 3:
                    qty_div = cols[0].find("div")
                    desc_td = cols[1]
                    if qty_div and desc_td:
                        # Inside the loop where you process each item
                        qty = qty_div.get_text(strip=True)
                        main_desc_div = desc_td.find("div", style=lambda s: s and "font-weight:700" in s)
                        if not main_desc_div:
                            continue
                        item_name = main_desc_div.get_text(strip=True)

                        # Extract special instructions more reliably
                        instructions_text = ""
                        instructions_div = desc_td.find("div", string=lambda text: text and "Instructions:" in text)
                        if not instructions_div:
                            # Try more general search inside the item block
                            for div in desc_td.find_all("div"):
                                if "Instructions:" in div.get_text():
                                    full_text = div.get_text(strip=True)
                                    instructions_text = re.sub(r"^\s*Instructions:\s*", "", full_text, flags=re.I)
                                    break
                        else:
                            full_text = instructions_div.get_text(strip=True)
                            instructions_text = re.sub(r"^\s*Instructions:\s*", "", full_text, flags=re.I)

                        # Sub-details in <ul><li>
                        details = []
                        ul = desc_td.find("ul")
                        if ul:
                            for li in ul.find_all("li"):
                                detail = li.get_text(strip=True)
                                if detail:
                                    details.append(f"- {detail}")

                        # Build item block
                        lines = [f"{qty} {item_name}"]
                        if instructions_text:
                            lines.append(f"Instructions: {instructions_text}")
                        if details:
                            lines.extend(details)

                        item_text = "\n".join(lines)
                        items.append(item_text)


            if not items:
                continue

            total_items = len(items)

            for idx, item in enumerate(items, start=1):
                label_lines = [
                    name_line,
                    order_date,
                    "-------------------------------------",
                    f"Item {idx} out of {total_items}",
                    item,
                    "-------------------------------------",
                    "Grubhub STO",
                    f"Deliver to: {delivery_name}"
                ]
                labels.append("\n".join(label_lines))

        except Exception as e:
            print(f"Error parsing block: {e}")
            continue

    return labels

def clear_table(table):
    for row in table.rows:
        for cell in row.cells:
            cell.text = ""

def set_cell_margins(cell, top=0, start=100, bottom=0, end=100):
    """
    Set cell margins using w:tcMar in twips (1 pt = 20 twips).
    start/end = left/right
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create <w:tcMar> element
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is not None:
        tcPr.remove(tcMar)
    tcMar = OxmlElement('w:tcMar')

    for side, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(value))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)

    tcPr.append(tcMar)


def fill_table(table: Table, label_texts):
    label_rows = [0, 2]  # Only these rows contain labels
    max_cols = len(table.rows[0].cells)
    total_cells = len(label_rows) * max_cols

    for idx, label in enumerate(label_texts):
        if idx >= total_cells:
            print(f"Warning: More labels ({len(label_texts)}) than cells ({total_cells}) in table.")
            break

        row = label_rows[idx // max_cols]
        col = idx % max_cols

        cell = table.rows[row].cells[col]
        set_cell_margins(cell, start=200, end=200)  # 200 twips = 10pt ≈ 3.5mm
        cell.text = ""  # Clear previous content

        lines = label.splitlines()

        for j, line in enumerate(lines):
            p = cell.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(0)  # Remove space after

            run = p.add_run(line)
            run.font.name = 'Arial'
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), 'Arial')

            if j == 0:
                run.bold = True
                run.font.size = Pt(14)
            elif line.lower().startswith("item") and "out of" in line.lower():
                run.bold = True
                run.font.size = Pt(10)
            elif line.strip() == "-------------------------------------":
                run.font.size = Pt(10)

                # Explicitly clear all indent
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.first_line_indent = Pt(0)

                # Optional: center the line
                p.alignment = 1  # 0 = left, 1 = center, 2 = right

            elif line.strip().startswith("-"):
                p.paragraph_format.left_indent = Pt(12)
                run.font.size = Pt(10)
            else:
                run.font.size = Pt(10)

def create_labels_doc(labels, title, template_path=TEMPLATE_PATH):
    if not labels:
        raise ValueError("No labels provided to generate the document.")

    doc = Document(template_path)

    base_table = doc.tables[0]
    clear_table(base_table)

    # Split labels into pages of 10
    chunks = [labels[i:i + LABELS_PER_PAGE] for i in range(0, len(labels), LABELS_PER_PAGE)]

    # Fill first table (template)
    fill_table(base_table, chunks[0])

    # For remaining pages
    for label_chunk in chunks[1:]:
        # Start from a clean copy of the *empty* base table
        empty_table_xml = copy.deepcopy(base_table._tbl)
        clear_table(Table(empty_table_xml, doc))  # Clear any leftover labels

        doc._body._body.append(empty_table_xml)
        new_table = Table(empty_table_xml, doc)
        fill_table(new_table, label_chunk)

    # Remove any trailing empty paragraphs between tables (optional cleanup)
    while doc.paragraphs and doc.paragraphs[-1].text.strip() == "":
        p = doc.paragraphs[-1]._element
        doc._body._body.remove(p)

    # Set document title property
    core_properties = doc.core_properties
    core_properties.title = title

    # Save the document to an in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Reset buffer position to the beginning
    print(f"Generated {len(labels)} labels in memory.")
    
    return buffer.getvalue()  # Return the bytes


st.title("Grubhub Label Generator")
uploaded_file = st.file_uploader("Upload .eml file containing orders", type=['eml'])
if uploaded_file:
    # Read the file contents
    file_contents = uploaded_file.read()
    html_content = convert_eml_to_html(file_contents)

    labels = extract_individual_labels_from_html(html_content)

    doc_bytes = create_labels_doc(labels, "Grubhub Labels")

    # Create download button with the document bytes
    st.download_button(
        label="Download Labels",
        data=doc_bytes,
        file_name="labels.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )